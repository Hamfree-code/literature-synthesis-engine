"""Editable DOCX deliverable (UPGRADE v3.1 — P6.5).

Renders a Markdown report into a corporately-styled Word document (python-docx)
so client teams can work the deliverable in Word: cover page, confidentiality
notice, document version-control table, auto TOC field, page numbers, and
consistent heading styles. The Markdown→DOCX conversion is intentionally simple
(headings, paragraphs, bullet lists, pipe tables, blockquotes) — enough for the
reports this engine emits, with no new heavy dependency.
"""

from __future__ import annotations

import re
from datetime import date


def _add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def _add_toc(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Right-click and 'Update Field' to build the table of contents."
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run._r.append(txt)
    run._r.append(fld3)


def _cover(document, title: str, subtitle: str, run_id: str, search_date: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document.add_paragraph()
    h = document.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Hams & Co. Research Division")
    r.bold = True
    r.font.size = __import__("docx").shared.Pt(20)

    t = document.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = __import__("docx").shared.Pt(26)

    s = document.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Evidence current as of {search_date}\nRun ID: {run_id}\nGenerated {date.today().isoformat()}"
    )

    document.add_paragraph()
    conf = document.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = conf.add_run("CONFIDENTIAL — for the intended recipient only.")
    cr.italic = True

    # Version control table.
    document.add_paragraph()
    vt = document.add_paragraph()
    vt.add_run("Document version control").bold = True
    table = document.add_table(rows=2, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Version", "Date", "Author", "Notes"]):
        hdr[i].text = label
    row = table.rows[1].cells
    row[0].text = "1.0"
    row[1].text = date.today().isoformat()
    row[2].text = "Literature Synthesis Engine"
    row[3].text = "Initial issue"

    document.add_page_break()


_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", line)) and "-" in line


def markdown_to_docx(
    md_body: str,
    out_path,
    *,
    title: str = "Research Report",
    subtitle: str = "",
    run_id: str = "n/a",
    search_date: str | None = None,
) -> str:
    """Convert a Markdown body into a styled .docx and return the path."""
    from pathlib import Path

    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    search_date = search_date or date.today().isoformat()
    document = docx.Document()

    section = document.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer_p)

    _cover(document, title, subtitle or title, run_id, search_date)

    document.add_heading("Table of Contents", level=1)
    _add_toc(document)
    document.add_page_break()

    lines = md_body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # Tables: header row + separator + body rows.
        if _TABLE_ROW.match(line) and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            body_rows = []
            j = i + 2
            while j < len(lines) and _TABLE_ROW.match(lines[j]):
                body_rows.append([c.strip() for c in lines[j].strip().strip("|").split("|")])
                j += 1
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Light List Accent 1"
            for k, cell in enumerate(table.rows[0].cells):
                cell.text = _strip_md(header[k] if k < len(header) else "")
            for br in body_rows:
                cells = table.add_row().cells
                for k in range(len(header)):
                    cells[k].text = _strip_md(br[k]) if k < len(br) else ""
            i = j
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            document.add_heading(_strip_md(line.lstrip("#").strip()), level=min(level, 4))
        elif line.lstrip().startswith(("- ", "* ")):
            document.add_paragraph(_strip_md(line.lstrip()[2:]), style="List Bullet")
        elif line.startswith(">"):
            p = document.add_paragraph(_strip_md(line.lstrip(">").strip()))
            p.runs[0].italic = True if p.runs else None
        else:
            document.add_paragraph(_strip_md(line))
        i += 1

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return str(out_path)


def _strip_md(text: str) -> str:
    """Strip inline markdown/HTML noise for Word (bold/italic/links/sup badges)."""
    text = re.sub(r"<sup>(.*?)</sup>", r" (\1)", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()
